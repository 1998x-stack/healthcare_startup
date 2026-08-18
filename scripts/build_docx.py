"""Build / update the project DOCX deliverables.

1. New  : 项目架构与技术方案.docx  (fresh technical architecture doc)
2. Edit : 基于ChatGPT的患者信息管理与医疗服务优化.docx
          (insert a 16. 技术架构概览（重构） section before the Prompt section)

Run with the docx skill's managed venv:
  /Users/x/.workbuddy/binaries/python/envs/default/bin/python scripts/build_docx.py
"""
import os
import shutil

import docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCK_DIR = os.path.join(ROOT, 'data', 'docx_out')
EXISTING = os.path.join(ROOT, '基于ChatGPT的患者信息管理与医疗服务优化.docx')
NEW_DOCX = os.path.join(ROOT, '项目架构与技术路线.docx')

CJK = 'PingFang SC'


def _set_cjk(run, name=CJK):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name)
    rfonts.set(qn('w:ascii'), 'Menlo' if name == 'Menlo' else name)
    rfonts.set(qn('w:hAnsi'), 'Menlo' if name == 'Menlo' else name)


def add_para(doc, text, style=None, bold=False, size=None, align=None):
    p = doc.add_paragraph(style=style)
    run_node = p.add_run(text)
    _set_cjk(run_node)
    run_node.bold = bold
    if size:
        run_node.font.size = size
    if align:
        p.alignment = align
    return p


def add_code(doc, text):
    for line in text.splitlines():
        p = doc.add_paragraph()
        run_node = p.add_run(line if line else ' ')
        _set_cjk(run_node, 'Menlo')
        run_node.font.size = docx.shared.Pt(9)
    return


def add_table(doc, headers, rows, width_auto=True):
    tb = doc.add_table(rows=1, cols=len(headers))
    tb.style = 'Light Grid Accent 1'
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tb.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        _set_cjk(r)
        r.bold = True
    for row in rows:
        cells = tb.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            _set_cjk(r)
    return tb


def build_new_docx(path):
    doc = Document()
    base = doc.styles['Normal']
    base.font.name = CJK
    base.element.rPr.rFonts.set(qn('w:eastAsia'), CJK)

    add_para(doc, '项目架构与技术路线', style='Title', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'healthcare_startup —— 工厂模式重构说明', style='Subtitle',
             align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, '一、项目概述', style='Heading 1')
    add_para(doc, '本工程包含两部分：Python 医疗数据采集（爬虫）和基于 ChatGPT 的微信小程序 '
                  '（MedGPT）。本方案说明 Python 侧三个爬虫（医院 / 药品 / 医患问答）经工厂模式重构后的'
                  '架构：共享一条多进程流水线，字段抽取委托给各自的解析器，新增站点只需注册新解析器。')

    add_para(doc, '二、目录结构', style='Heading 1')
    add_code(doc, """healthcare_startup
├── scrapers/            # 爬虫包（工厂模式）
│   ├── __init__.py      # create_scraper / ScraperFactory
│   ├── base.py          # BaseScraper：流水线 + 多进程
│   ├── parsers.py       # Hospital / Medicine / QA 解析器
│   ├── scrapers.py      # 三种具体 Scraper
│   ├── factory.py       # ScraperFactory.create(site)
│   └── urls.py          # URL 构建
├── libs/                # utils / log / tools
├── config.py            # 集中配置
├── run.py               # CLI 入口
├── MedGPT/              # 微信小程序
├── docs/                # 文档
├── data/                # 采集数据
└── tests/               # 离线 pytest""")

    add_para(doc, '三、工厂模式架构', style='Heading 1')
    add_para(doc, '3.1 设计模式：工厂 + 解析器', style='Heading 2')
    add_para(doc, 'ScraperFactory.create(site) 根据站点名返回具体的 Scraper（BaseScraper 子类）。'
                  '每个 Scraper 只负责 build_tasks()，字段抽取委托给绑定的解析器。'
                  '新增站点只需新增解析器并注册到注册表。')
    add_para(doc, '3.2 组件职责', style='Heading 2')
    add_table(doc,
              ['文件', '职责', '对外接口'],
              [['scrapers/base.py', '共享流水线 + 多进程编排', 'BaseScraper.run() -> dict'],
               ['scrapers/parsers.py', '字段抽取', 'parse(soup, url, fetch, logger)'],
               ['scrapers/scrapers.py', '三种具体 Scraper', 'Hospital/Medicine/QAScraper'],
               ['scrapers/factory.py', '工厂', 'create(site)'],
               ['config.py', '集中配置', 'SITE_NAMES / SAVE_FOLDERS'],
               ['run.py', '命令行', 'python run.py site']])
    add_para(doc, '3.3 站点映射', style='Heading 2')
    add_table(doc,
              ['site', 'Scraper', 'Parser', '输出目录'],
              [['hospital', 'HospitalScraper', 'HospitalParser', 'data/hospital/'],
               ['medicine', 'MedicineScraper', 'MedicineParser', 'data/medician/'],
               ['qa', 'QAScraper', 'QAParser', 'data/ques_ans/']])

    add_para(doc, '四、多进程流水线', style='Heading 1')
    add_para(doc, 'BaseScraper.run() 使用 multiprocessing.JoinableQueue（修复了旧代码中 queue.Queue '
                  '导致的 “cannot pickle _thread.lock” 崩溃）。工作函数为模块级 _worker，'
                  '可跨进程按引用序列化；解析器通过注册表按名字解析，fetch 默认 utils.get_soup。')

    add_para(doc, '五、运行指南', style='Heading 1')
    add_code(doc, "python run.py hospital [--processes 8]\n"
                  "python run.py medicine [--processes 8] [--count 100000]\n"
                  "python run.py qa   [--processes 8] [--count 1000]")

    add_para(doc, '六、测试与验证', style='Heading 1')
    add_para(doc, 'pytest tests/ -v 全部离线运行（工厂、URL、BaseScraper、解析器 fixture、CLI）；'
                  'scripts/smoke_integration.py 端到端多进程冒烟。')

    add_para(doc, '七、已知限制', style='Heading 1')
    add_para(doc, 'data/hospital/东城/北京协和医院/ 下 55 个科室目录名因历史多次乱码转换已不可用启发式'
                  '恢复，保留为 renamed-* 避免误改数据。')

    doc.save(path)
    print('wrote', path)


def update_existing(path, out_path):
    doc = Document(path)
    target = None
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1' and p.text.strip() == 'Prompt':
            target = p
            break
    if target is None:
        print('WARN: Prompt heading not found; appending at end')
    # Build the new paragraphs to insert.
    new_items = []
    new_items.append(('H1', '16. 技术架构概览（重构）'))
    new_items.append(('P', '为保证本文项目的可靠采集，Python 侧三个爬虫已重构为工厂模式（scrapers/ 包）：'
                           'ScraperFactory.create(site) 返回绑定解析器的具体 Scraper，共享 BaseScraper '
                           '的多进程流水线（JoinableQueue）。'))
    new_items.append(('H2', '16.1 站点映射'))
    new_items.append(('P', 'hospital→HospitalScraper/HospitalParser→data/hospital/；'
                           'medicine→MedicineScraper/MedicineParser→data/medician/；'
                           'qa→QAScraper/QAParser→data/ques_ans/。'))
    new_items.append(('H2', '16.2 运行'))
    new_items.append(('P', 'python run.py hospital|medicine|qa [--processes N]。详见 README 与 docs/architecture.md。'))

    # Build a table before Prompt too
    anchor = target._element
    for kind, text in new_items:
        p = doc.add_paragraph(style='Heading 1' if kind == 'H1' else
                              ('Heading 2' if kind == 'H2' else None))
        r = p.add_run(text)
        _set_cjk(r)
        if kind == 'H1':
            r.bold = True
        anchor.addprevious(p._element)

    # Component mapping table inserted before Prompt
    tb = doc.add_table(rows=1, cols=4)
    tb.style = 'Normal Table'
    _add_table_borders(tb)
    headers = ['site', 'Scraper', 'Parser', '输出']
    for i, h in enumerate(headers):
        cell = tb.rows[0].cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        _set_cjk(r)
        r.bold = True
    rows = [['hospital', 'HospitalScraper', 'HospitalParser', 'data/hospital/'],
            ['medicine', 'MedicineScraper', 'MedicineParser', 'data/medician/'],
            ['qa', 'QAScraper', 'QAParser', 'data/ques_ans/']]
    for row in rows:
        cells = tb.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run(val)
            _set_cjk(r)
    anchor.addprevious(tb._element)

    doc.save(out_path)
    print('wrote', out_path)


def _add_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = tbl_pr.makeelement(qn('w:' + edge), {})
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tbl_pr.append(borders)


def main():
    os.makedirs(DOCK_DIR, exist_ok=True)
    build_new = True
    build_existing = True
    if build_new:
        build_new_docx(NEW_DOCX)
    if build_existing:
        update_existing(EXISTING, EXISTING)
    print('done')


if __name__ == '__main__':
    main()